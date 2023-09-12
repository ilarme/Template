#!/usr/bin/python3
# -*- coding: utf-8 -*-
import copy
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import sys
import psycopg2

def is_valid_utf_8_symbol(c):
    cd = ord(c)
    if cd not in range(0x0, 0x20) and cd != 0x7f:
        return True
    return False


def create_docx(program_id):
    def create_connection(db_name, db_user, db_password, db_host, db_port):
        connection = None
        try:
            connection = psycopg2.connect(
                database=db_name,
                user=db_user,
                password=db_password,
                host=db_host,
                port=db_port,
            )
            print("Connection to PostgreSQL DB successful")
        except psycopg2.OperationalError as e:
            print(f"The error '{e}' occurred")
        return connection

    connection = create_connection(
        sys.argv[1],  sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5]
    )
    cursor = connection.cursor()

    # Входные данные
    cursor.execute(
        "SELECT university_name, program_name, specialization_name "
        "FROM datace.view_university_rating_separate WHERE program_id = %s",
        (program_id,))
    info = cursor.fetchone()
    cursor.execute("SELECT expert_name FROM datace.view_expert_form_word_doc WHERE program_id = (%s)", (program_id,))
    experts = cursor.fetchall()
    cursor.execute(
        "SELECT expert_comment_1, expert_comment_2, expert_comment_3, expert_comment_4 "
        "FROM datace.view_expert_form_word_doc WHERE program_id = (%s)",
        (program_id,))
    m = cursor.fetchall()
    l0 = [list(n) for n in m]
    for i in range(0, len(l0)):
        k = 1
        for j in range(0, len(l0[i])):
            if l0[i][j] is None:
                continue
            l0[i][j] = "Комментарий № " + str(k) + ": " + l0[i][j].replace(chr(10), '') + ""
            k += 1
    l = ["Эксперт " + experts[i][0] + '\n' + "\n".join(filter(None, l0[i])) for i in range(0, len(l0)) if len(list(filter(None, l0[i]))) > 0]
    rec1 = "\n\n".join(filter(None, l))
    cursor.execute(
        "SELECT expert_comment_5, expert_comment_6, expert_comment_7, expert_comment_8, "
        "expert_comment_9 FROM datace.view_expert_form_word_doc WHERE program_id = %s",
        (program_id,))
    m = cursor.fetchall()
    l0 = [list(n) for n in m]
    for i in range(0, len(l0)):
        k = 1
        for j in range(0, len(l0[i])):
            if l0[i][j] is None: continue
            l0[i][j] = "Комментарий № " + str(k) + ": " + l0[i][j].replace(chr(10), '') + ""
            k += 1
    l = ["Эксперт " + experts[i][0] + '\n' + "\n".join(filter(None, l0[i])) for i in range(0, len(l0)) if len(list(filter(None, l0[i]))) > 0]
    rec2 = "\n\n".join(filter(None, l))
    cursor.execute(
        "SELECT expert_comment_10, expert_comment_11, expert_comment_12, comment_13 "
        "FROM datace.view_expert_form_word_doc WHERE program_id = %s",
        (program_id,))
    m = cursor.fetchall()
    l0 = [list(n) for n in m]
    for i in range(0, len(l0)):
        k = 1
        for j in range(0, len(l0[i])):
            if l0[i][j] is None: continue
            l0[i][j] = "Комментарий № " + str(k) + ": " + l0[i][j].replace(chr(10), '') + ""
            k += 1
    l = ["Эксперт " + experts[i][0] + '\n' + "\n".join(filter(None, l0[i])) for i in range(0, len(l0)) if len(list(filter(None, l0[i]))) > 0]
    rec3 = "\n\n".join(filter(None, l))
    recomendations = (rec1, rec2, rec3)
    cursor.execute("SELECT expert_comment_3 FROM datace.view_expert_form_word_doc WHERE program_id = %s", (program_id,))
    demand = ""
    cursor.execute("SELECT final_conclusion FROM datace.view_expert_form_word_doc WHERE program_id = %s", (program_id,))
    c = cursor.fetchall()
    accept = reject = 0
    for s in c:
        if s[0].lstrip() == 'Соответствует':
            accept += 1
        if s[0].lstrip() == 'Не соответсвует':
            reject += 1
    if accept > reject:
        conclusion = 'Соответствует'
    if accept < reject:
        conclusion = 'Требуется доработка'
    if accept == reject:
        conclusion = 'Спорно'
    cursor.execute("SELECT improvement_recommendations FROM datace.view_expert_form_word_doc WHERE program_id = %s",
                   (program_id,))
    m = cursor.fetchall()
    l0 = [list(n) for n in m]
    for i in range(0, len(l0)):
        if l0[i][0] is None:
            continue
        l0[i][0] = "" + l0[i][0].replace(chr(10), '') + ""
    l = ["Эксперт " + experts[i][0] + '\n' + "\n".join(filter(None, l0[i])) for i in range(0, len(l0))]
    final_conclusion = "\n\n".join(filter(None, l))


    document = Document()

    # горизонтальная линия
    def insertHR(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
                                  'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                  'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                  'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                  'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                  'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                                  'w:pPrChange'
                                  )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def page_break():
        """Метод Document.add_page_break()"""
        from docx.enum.text import WD_BREAK
        p = document.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph()
        return p


    # Определяем стиль заголовка
    style_heading = document.styles.add_style('Style_heading', WD_STYLE_TYPE.PARAGRAPH)
    style_heading.font.name = 'Times New Roman'
    style_heading.font.size = Pt(11)
    style_heading.font.bold = True

    style_heading1 = document.styles.add_style('Style_heading1', WD_STYLE_TYPE.PARAGRAPH)
    style_heading1.font.name = 'Times New Roman'
    style_heading1.font.size = Pt(10)
    style_heading1.font.bold = True

    style1 = document.styles.add_style('Normal1', WD_STYLE_TYPE.PARAGRAPH)
    style1.font.name = 'Times New Roman'
    style1.font.size = Pt(10)

    style_header = document.styles.add_style('Style_header', WD_STYLE_TYPE.PARAGRAPH)
    style_header.font.name = 'Calibri Light'
    style_header.font.size = Pt(9)

    style_header1 = document.styles.add_style('Style_header1', WD_STYLE_TYPE.PARAGRAPH)
    style_header1.font.name = 'Calibri Light'
    style_header1.font.size = Pt(12)
    style_header1.font.bold = True

    # Определяем стиль текста
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    section1 = document.sections[0]
    section1.page_width = Cm(21)
    section1.page_height = Cm(29.7)

    section1.left_margin = Cm(2.54)
    section1.right_margin = Cm(1.42)
    section1.top_margin = Cm(2.54)
    section1.bottom_margin = Cm(1.25)

    run = document.add_paragraph(
        'Экспертное заключение по оценке дополнительной профессиональной программы\n'
        'профессиональной переподготовки (ДПП ПП) '
        'или программ обучения по модулям ИТ-профиля\n'
        'в пределах основной образовательной профессиональной программы высшего '
        'образования \n(модуль ИТ-профиля ОПОП ВО)\n', style=style_heading)
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    header = section1.header

    table = header.add_table(1, 2, section1.page_width)
    par = table.columns[0].cells[0].add_paragraph()
    r = par.add_run()
    r.add_picture('img1.png', width=Cm(7.5))
    par = table.columns[1].cells[0].add_paragraph("АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ", style=style_header)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    par.paragraph_format.space_after = Pt(0)

    par = table.columns[1].cells[0].add_paragraph("ЦИФРОВАЯ ЭКОНОМИКА", style=style_header1)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)

    insertHR(header.add_paragraph(style=style_header1))

    footer = section1.footer
    insertHR(footer.add_paragraph(style=style_header1))
    table = footer.add_table(1, 2, section1.page_width)
    par = table.columns[0].cells[0].add_paragraph()
    r = par.add_run()
    r.add_picture('img2.png', width=Cm(1.5))
    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    par = table.columns[1].cells[0].add_paragraph("ПРОЕКТ «КОНТРОЛЬ КАЧЕСТВА\nОБРАЗОВАНИЯ»", style=style_header1)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    table.columns[0].cells[0].width = Cm(12)

    data = str(datetime.now().date())
    if info is None:
        print("No reviews found")
        return False
    strings = list(info)
    for i in range(0, len(info)):
        strings[i] = "".join(c for c in strings[i] if is_valid_utf_8_symbol(c))
    items = (
        (1, 'Дата составления', data),
        (2, 'Наименование вуза', strings[0].lstrip()),
        (3, 'Наименование дополнительной профессиональной программы профессиональной переподготовки', strings[1]),
        (4, 'Область цифровых компетенций', strings[2]),
    )
    # таблица
    table = document.add_table(0, 3)

    for item in items:
        cells = table.add_row().cells
        cells[0].text = str(item[0])
        cells[0].paragraphs[0].style = style_heading
        cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[1].text = item[1]
        cells[1].paragraphs[0].style = style_heading
        cells[2].text = item[2]
        cells[2].paragraphs[0].style = style

    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for column in table.columns:
        for cell in column.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(8)
            cell.paragraphs[0].paragraph_format.space_before = Pt(8)

    for cell in table.columns[0].cells:
        cell.width = Cm(0.93)
    for cell in table.columns[1].cells:
        cell.width = Cm(7.06)
    for cell in table.columns[2].cells:
        cell.width = Cm(8)

    section2 = document.add_section(WD_SECTION_START.NEW_PAGE)
    new_width, new_height = Cm(29.7), Cm(21)
    section2.orientation = WD_ORIENTATION.LANDSCAPE
    section2.page_width = new_width
    section2.page_height = new_height

    header = section2.header
    header.is_linked_to_previous = False

    footer = section2.footer
    footer.is_linked_to_previous = False

    run = document.add_paragraph(
        'Протокол оценки соответствия дополнительной профессиональной программы профессиональной '
        'переподготовки ИТ-профиля или программы обучения по модулям ИТ-профиля в пределах '
        'основной образовательной профессиональной программы высшего образования запросам '
        'приоритетных отраслей экономики, в том числе ИТ-отрасли в рамках реализации '
        'проекта\n «Цифровые кафедры»', style=style_heading)
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    section2 = document.sections[0]
    section2.left_margin = Cm(1.25)
    section2.right_margin = Cm(1.45)
    section2.top_margin = Cm(1.5)
    section2.bottom_margin = Cm(1.75)
    if conclusion == "Спорно" or conclusion == "Не соответствует":
        table = document.add_table(1, 4)
    else:
        table = document.add_table(1, 3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    items = ("Критерии соответствия",
             "Рекомендации", "Отработка замечаний экспертов (заполняет вуз)")
    for i in range (1, len(table.columns)):
        table.rows[0].cells[i].text = items[i - 1]
        table.rows[0].cells[i].paragraphs[0].style = style_heading1
        table.rows[0].cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    items = (
        ("1.", "Отраслевая компонента:",
         "Соответствие образовательной программы требованиям отрасли, с учетом аспектов:\n"
         "  •  Прикладного характера программ\n"
         "  •  Проектно-ориентированного характера программ\n"
         "  •  Востребованности предлагаемых знаний и навыков по основному направлению обучения (для данной специальности)\n"
         "  •  Иные отраслевые блоки",
         recomendations[0]
         ),
        ("2.", "ИТ компонента:",
         "Соответствие образовательной программы требованиям отрасли с точки зрения ИТ с учетом аспектов:\n"
         "  •  Актуальности знаний, навыков, технологий\n"
         "  •  Наличие практической составляющей для получения практического опыта\n"
         "  •  Импортозамещение и отечественного программного обеспечения\n"
         "  •  Информационной и кибер безопасности\n"
         "  •  Опционально: Нацеленность на экспорт создаваемой продукции",
         recomendations[1]
         ),
        ("3.", "Образовательная компонента:",
         "Соответствие образовательной программы требованиям отрасли с точки зрения опыта педагогов в части:\n"
         "  •  Реализации отраслевых проектов\n"
         "  •  Взаимодействия с отраслью\n"
         "  •  Вариативность форм и методов обучения\n"
         "  •  Иные отраслевые блоки",
         recomendations[2]
         )
    )

    for item in items:
        cells = table.add_row().cells
        cells[0].text = item[0]
        cells[0].paragraphs[0].style = style1
        cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[1].text = item[1]
        cells[1].paragraphs[0].style = style_heading1
        par = cells[1].add_paragraph(item[2])
        par.style = style1
        cells[2].text = item[3]
        cells[2].paragraphs[0].style = style1

    for column in table.columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.space_before = Pt(6)

    for cell in table.columns[0].cells:
        cell.width = Cm(1.56)
    for cell in table.columns[1].cells:
        cell.width = Cm(7.17)
    for cell in table.columns[2].cells:
        cell.width = Cm(9.6)

    table = document.add_table(0, 3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    items = (
        ("4.", "Востребованность специалиста после прохождения программы на рынке труда через 3 года.", demand),
        ("5.", "Заключение о соответствии программы требованиям отрасли", conclusion),
        ("6.", "Итоговое заключение о соответствии или несоответствии дополнительной профессиональной программы "
               "профессиональной переподготовки критериям отбора (соответствует/не соответствует) и основные "
               "рекомендации по улучшения программы.", final_conclusion
         )
    )
    for item in items:
        cells = table.add_row().cells
        cells[0].text = item[0]
        cells[0].paragraphs[0].style = style1
        cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[1].text = item[1]
        cells[1].paragraphs[0].style = style_heading1
        cells[2].text = item[2]
        cells[2].paragraphs[0].style = style_heading1
    table.rows[2].cells[1].paragraphs[0].style = style1
    table.rows[2].cells[2].paragraphs[0].style = style1
    table.rows[0].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.rows[1].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for column in table.columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.space_before = Pt(6)
    for cell in table.columns[0].cells:
        cell.width = Cm(1.56)
    for cell in table.columns[1].cells:
        cell.width = Cm(7.17)
    for cell in table.columns[2].cells:
        cell.width = Cm(18.25)
    # Сохраняем документ
    # document.save('Экспертное заключение по программе ' + str(program_id) + ' ' + str(datetime.now().date()) + '.docx')
    # print("Document created successfully")
    # return True
    return document

if __name__ == '__main__':
    for program_id in range(9, 10):
        d = create_docx(program_id)
        d.save(
            'Экспертное заключение по программе ' + str(program_id) + ' ' + str(datetime.now().date()) + '.docx')
